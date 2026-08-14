from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
REFERENCE = ROOT / "deliverables" / "麻将小程序游戏规则需求规格_v0.1.docx"
SOURCE = ROOT / "docs" / "GAME_RULES.md"
OUTPUT = ROOT / "deliverables" / "麻将小程序游戏规则需求规格_v0.2.docx"

RULE_RE = re.compile(
    r"^- \*\*(?P<code>[RP]\d+)(?:（待确认）)?\*\*\s+(?P<text>.+)$"
)


def text_nodes(paragraph_element):
    return list(paragraph_element.iter(qn("w:t")))


def clone_single(pattern, text: str):
    paragraph = deepcopy(pattern)
    nodes = text_nodes(paragraph)
    if not nodes:
        raise ValueError("Paragraph pattern has no text node")
    nodes[0].text = text
    for node in nodes[1:]:
        node.text = ""
    return paragraph


def clone_split(pattern, first: str, second: str):
    paragraph = deepcopy(pattern)
    nodes = text_nodes(paragraph)
    if len(nodes) < 2:
        raise ValueError("Split paragraph pattern has fewer than two text nodes")
    nodes[0].text = first
    nodes[-1].text = second
    for node in nodes[1:-1]:
        node.text = ""
    return paragraph


def parse_source():
    title = None
    version = None
    boundary = None
    sections: list[tuple[str, list[tuple[str, str]]]] = []
    current_rules: list[tuple[str, str]] | None = None
    footer_note = None

    for raw_line in SOURCE.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            title = line[2:].strip()
            continue
        if line.startswith("版本 v"):
            version = line
            continue
        if line.startswith("> "):
            boundary = line[2:].strip()
            continue
        if line.startswith("## "):
            current_rules = []
            sections.append((line[3:].strip(), current_rules))
            continue
        match = RULE_RE.match(line)
        if match:
            if current_rules is None:
                raise ValueError(f"Rule outside section: {line}")
            current_rules.append((match.group("code"), match.group("text")))
            continue
        if line.startswith("版本控制："):
            footer_note = line
            continue
        raise ValueError(f"Unsupported source line: {line}")

    if not all((title, version, boundary, footer_note)):
        raise ValueError("Source is missing title, version, boundary, or footer note")
    return title, version, boundary, sections, footer_note


def build():
    title, version, boundary, sections, footer_note = parse_source()
    document = Document(REFERENCE)
    paragraphs = document.paragraphs
    patterns = {
        "label": deepcopy(paragraphs[0]._p),
        "title": deepcopy(paragraphs[1]._p),
        "version": deepcopy(paragraphs[2]._p),
        "callout": deepcopy(paragraphs[3]._p),
        "heading": deepcopy(paragraphs[4]._p),
        "rule": deepcopy(paragraphs[5]._p),
        "page_break": deepcopy(paragraphs[100]._p),
        "pending": deepcopy(paragraphs[102]._p),
        "footer_note": deepcopy(paragraphs[115]._p),
    }

    body = document._element.body
    section_properties = body.sectPr
    for child in list(body):
        if child is not section_properties:
            body.remove(child)

    body.insert(-1, clone_single(patterns["label"], "规则需求规格"))
    body.insert(-1, clone_single(patterns["title"], title))
    body.insert(-1, clone_single(patterns["version"], version))
    callout_prefix = "文档边界："
    callout_body = boundary.removeprefix(callout_prefix)
    body.insert(-1, clone_split(patterns["callout"], callout_prefix, callout_body))

    for heading, rules in sections:
        is_pending = "待确认事项" in heading
        if is_pending:
            body.insert(-1, deepcopy(patterns["page_break"]))
        body.insert(-1, clone_single(patterns["heading"], heading))
        rule_pattern = patterns["pending"] if is_pending else patterns["rule"]
        for code, text in rules:
            body.insert(-1, clone_split(rule_pattern, code, text))

    footer_prefix = "版本控制："
    footer_body = footer_note.removeprefix(footer_prefix)
    body.insert(-1, clone_split(patterns["footer_note"], footer_prefix, footer_body))

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            nodes = text_nodes(paragraph._p)
            if nodes:
                nodes[0].text = paragraph.text.replace("v0.1", "v0.2")
                for node in nodes[1:]:
                    node.text = ""

    document.core_properties.title = title
    document.core_properties.subject = "麻将小程序已确认规则固化稿 v0.2"
    document.core_properties.comments = "Generated from docs/GAME_RULES.md; v0.1 retained unchanged."
    document.save(OUTPUT)


def verify():
    source_rules = []
    for line in SOURCE.read_text(encoding="utf-8").splitlines():
        match = RULE_RE.match(line.strip())
        if match:
            source_rules.append((match.group("code"), match.group("text")))

    document = Document(OUTPUT)
    output_rules = []
    for paragraph in document.paragraphs:
        if paragraph.style.name not in {"Rule", "Pending Rule"}:
            continue
        code, separator, text = paragraph.text.partition("\t")
        if not separator:
            raise AssertionError(f"Rule paragraph has no tab: {paragraph.text}")
        output_rules.append((code, text))

    if output_rules != source_rules:
        raise AssertionError("DOCX rule list differs from Markdown source")

    codes = [code for code, _ in output_rules]
    if len(codes) != len(set(codes)):
        raise AssertionError("Duplicate rule or pending code found")
    if any("待补充" in text or "TODO" in text for _, text in output_rules):
        raise AssertionError("Placeholder text found")
    if not OUTPUT.exists() or OUTPUT.stat().st_size < 10_000:
        raise AssertionError("Output DOCX is missing or unexpectedly small")

    confirmed = sum(code.startswith("R") for code in codes)
    pending = sum(code.startswith("P") for code in codes)
    print(f"Created: {OUTPUT}")
    print(f"Confirmed rules: {confirmed}")
    print(f"Pending items: {pending}")
    print(f"Size: {OUTPUT.stat().st_size} bytes")


if __name__ == "__main__":
    build()
    verify()
